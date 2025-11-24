"""Export module for converting essays to various formats."""

import logging
import markdown2
from pathlib import Path
from typing import Optional

# Import optional dependencies with error handling
try:
    from weasyprint import HTML, CSS
    WEASYPRINT_AVAILABLE = True
except (ImportError, OSError):
    WEASYPRINT_AVAILABLE = False

try:
    from docx import Document
    from docx.shared import Pt
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

logger = logging.getLogger(__name__)

class Exporter:
    """Handles exporting essays to different formats."""

    def __init__(self):
        """Initialize exporter."""
        pass

    def export(self, text: str, format: str, output_path: str) -> bool:
        """
        Export text to the specified format.

        Args:
            text: The markdown text to export.
            format: Target format (pdf, docx, html, markdown).
            output_path: Destination file path.

        Returns:
            True if successful, False otherwise.
        """
        format = format.lower()
        path = Path(output_path)
        
        # Ensure parent directory exists
        path.parent.mkdir(parents=True, exist_ok=True)

        try:
            if format == "html":
                return self.to_html(text, path)
            elif format == "pdf":
                return self.to_pdf(text, path)
            elif format == "docx":
                return self.to_docx(text, path)
            elif format == "markdown" or format == "md":
                path.write_text(text)
                return True
            else:
                logger.error(f"Unsupported format: {format}")
                return False
        except Exception as e:
            logger.error(f"Export failed: {e}")
            return False

    def to_html(self, text: str, output_path: Path) -> bool:
        """Convert markdown to HTML."""
        html_content = markdown2.markdown(text, extras=["fenced-code-blocks", "tables"])
        
        # Add basic styling
        styled_html = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="utf-8">
            <style>
                body {{
                    font-family: 'Georgia', serif;
                    line-height: 1.6;
                    max-width: 800px;
                    margin: 40px auto;
                    padding: 20px;
                    color: #333;
                }}
                h1, h2, h3 {{ color: #2c3e50; }}
                code {{ background: #f4f4f4; padding: 2px 5px; border-radius: 3px; }}
                pre {{ background: #f4f4f4; padding: 15px; border-radius: 5px; overflow-x: auto; }}
                blockquote {{ border-left: 4px solid #ccc; margin: 0; padding-left: 16px; color: #666; }}
            </style>
        </head>
        <body>
            {html_content}
        </body>
        </html>
        """
        
        output_path.write_text(styled_html)
        return True

    def to_pdf(self, text: str, output_path: Path) -> bool:
        """Convert markdown to PDF via HTML."""
        if not WEASYPRINT_AVAILABLE:
            logger.error("WeasyPrint not installed. Cannot export to PDF.")
            return False

        # First convert to HTML
        html_content = markdown2.markdown(text, extras=["fenced-code-blocks", "tables"])
        
        # Professional PDF styling
        css = CSS(string="""
            @page {
                size: Letter;
                margin: 1in;
            }
            body {
                font-family: 'Times New Roman', serif;
                font-size: 12pt;
                line-height: 2.0; /* Double spacing for academic look */
            }
            h1 { font-size: 18pt; text-align: center; margin-bottom: 24pt; }
            h2 { font-size: 14pt; margin-top: 18pt; }
            p { margin-bottom: 12pt; text-align: justify; }
            blockquote { margin-left: 0.5in; font-style: italic; }
        """)

        HTML(string=html_content).write_pdf(output_path, stylesheets=[css])
        return True

    def to_docx(self, text: str, output_path: Path) -> bool:
        """Convert markdown to DOCX (simple conversion)."""
        if not DOCX_AVAILABLE:
            logger.error("python-docx not installed. Cannot export to DOCX.")
            return False

        doc = Document()
        
        # Set default font to Times New Roman 12pt
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(12)

        # Simple line-by-line parser
        # Note: This is a basic implementation. Complex markdown (nested lists, etc.) 
        # might not render perfectly, but it covers the basics.
        lines = text.split('\n')
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
                
            if line.startswith('# '):
                doc.add_heading(line[2:], level=1)
            elif line.startswith('## '):
                doc.add_heading(line[3:], level=2)
            elif line.startswith('### '):
                doc.add_heading(line[4:], level=3)
            elif line.startswith('- ') or line.startswith('* '):
                doc.add_paragraph(line[2:], style='List Bullet')
            elif line.startswith('1. '):
                # Simple check for numbered lists
                doc.add_paragraph(line[3:], style='List Number')
            elif line.startswith('> '):
                doc.add_paragraph(line[2:], style='Quote')
            else:
                doc.add_paragraph(line)

        doc.save(output_path)
        return True
